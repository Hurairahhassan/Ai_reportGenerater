import streamlit as st
import pandas as pd
from io import BytesIO
import pdfkit


st.title('Word File Generator')

col1, col2 = st.columns(2) 

df1 = df2 = None

with col1:
    # File uploader for the first timesheet report file
    uploaded_file1 = st.file_uploader("Choose the Timesheet report file", type=["csv", "xlsx"], key="file1")

    if uploaded_file1 is not None:
        try:
            if uploaded_file1.name.endswith('.csv'):
                df1 = pd.read_csv(uploaded_file1)
            elif uploaded_file1.name.endswith('.xlsx'):
                df1 = pd.read_excel(uploaded_file1)
            
            st.write("Timesheet report file uploaded successfully!")
            st.write(df1)
        except Exception as e:
            st.error(f"Error while loading the timesheet report: {e}")
    else:
        st.write("Please upload the timesheet report CSV or Excel file.")

with col2:
    # File uploader for the second timesheet report file
    uploaded_file2 = st.file_uploader("Choose the Funding Status file", type=["csv", "xlsx"], key="file2")

    if uploaded_file2 is not None:
        try:
            if uploaded_file2.name.endswith('.csv'):
                df2 = pd.read_csv(uploaded_file2)
            elif uploaded_file2.name.endswith('.xlsx'):
                df2 = pd.read_excel(uploaded_file2)
            
            st.write("Funding Status file uploaded successfully!")
            st.write(df2)
        except Exception as e:
            st.error(f"Error while loading the Funding Status file: {e}")
    else:
        st.write("Please upload the second Funding Status CSV or Excel file.")



if df1 is not None and df2 is not None:

    if st.button('Start Processing'):
        
        from docx.api import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Pt
        # from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        import pandas as pd
        import google.generativeai as genai
        import json

        count=0
        while count<5 :
        
            count=count+1
            try:
                def get_json_from_excel(df):
                    
                    # # Read Excel files
                    # def read_excel(file_path):
                    #     df = pd.read_excel(file_path)
                    #     return df

                    # Process your attachment
                    excel_data = df

                    combined_data = excel_data.to_string() + '''
                    \n please read that xlsx file data and convert that data into this json format whenever I give you xlsx file \n this is the json format which I want: \n [
                        {
                            "First Table": [
                                {"DAY OF SERVICE":"Monday","DATE OF SERVICE PROVIDED":"2024-4-1","SHIFT START":"9:00am","SHIFT END":"11:00am","HOURS":"2"},
                                {"DAY OF SERVICE":"Tuesday","DATE OF SERVICE PROVIDED":"2024-4-2","SHIFT START":null,"SHIFT END":null,"HOURS":"0"},
                                ...
                            ],
                            "Second Table": {"Total Regular Hours":"18","Back-up Hours":"18","Grand Total":"36"},
                            "Third Table": {"PAS": "18","PAS-BACKUP": "18","PROGRAM MANAGEMENT": "32", "ADMINISTRATIVE SUPPORT": "32","Total Labor Hours": "100"}
                        }
                    ] 
                    \n remember I am going to fit your json response in my python script only give me output in above specific json format (with same data types and date format) whenever I hit this prompt with excel file and there is a chance every time excel file does not follow this specific format you need to analyze that excel file and then convert that into that above json format
                    '''

                    genai.configure(api_key='AIzaSyAbxQrYBGdfm8DDNRXYrMSH8rWPRoeraCk')
                    model = genai.GenerativeModel('gemini-pro')
                    response = model.generate_content(combined_data)
                    # print(response.text)
                    return json.loads(response.text)

                def get_json_from_excel2(df2):
                    # Read Excel files
                    # def read_excel(file_path):
                    #     df = pd.read_excel(file_path)
                    #     return df

                    # Process your attachment
                    excel_data = df2

                    combined_data = excel_data.to_string() + '''
                    \n please read that xlsx file data and convert that data into this json format whenever I give you xlsx file \n this is the json format which I want: \n 
                        [
                            {"Service Period":"Jun 2023","Invoice Number":"3064","Labor":"131","Amount Claimed - Current Period":"$9267.19","Amount Remaining":"$104012.81","Total Amount Claimed":"$9267.19","Percentage Utilized":"8%"},...
                            {"Service Period":"Jun 2023","Invoice Number":"3271","Labor":"107.75","Amount Claimed - Current Period":"$7470.03","Amount Remaining":"$65613.38","Total Amount Claimed":"$158520.7","Percentage Utilized":"71%"},
                            {"Service Period":"Sep 2023","Invoice Number":"3298","Labor":"100","Amount Claimed - Current Period":"$6914.12","Amount Remaining":"$58699.26","Total Amount Claimed":"$165434.82","Percentage Utilized":"74%"},
                            {"Service Period":"Aug 2023","Invoice Number":null,"Labor":null,"Amount Claimed - Current Period":null,"Amount Remaining":"$58699.26","Total Amount Claimed":"$165434.82","Percentage Utilized":"74%"}
                        ]
                    \n remember I am going to fit your json response in my python script only give me output in above specific json format (with same data types, same rounded percentage values and same date format) whenever I hit this prompt with excel file and there is a chance every time excel file does not follow this specific format you need to analyze that excel file and then convert that into that above json format
                    '''

                    genai.configure(api_key='AIzaSyAbxQrYBGdfm8DDNRXYrMSH8rWPRoeraCk')
                    model = genai.GenerativeModel('gemini-pro')
                    response = model.generate_content(combined_data)
                    # print(response.text)
                    return json.loads(response.text)
                
                count2=0
                while count2<3:
                    count2=count2+1 
                    try:
                        data1= get_json_from_excel(df1)
                        break
                    except Exception as e:
                        # print("Error while loading data 1",e)
                        pass

                count3=0
                while count3<3:
                    count3=count3+1
                    try:
                        data2= get_json_from_excel2(df2)
                        if data2[0][1]["Service Period"].endswith("00:00:00"):
                            raise ValueError("Gemini Response is wrong")
                        break
                    except Exception as e:
                        # print("Error while loading data 2",e)
                        pass
                        

                if len(data1[0]["First Table"])==22:
                    document= Document("test.docx")

                if len(data1[0]["First Table"])==21:
                    document= Document("test2.docx")

                last_table= document.tables[-1]
                First_table= document.tables[0]
                Second_table= document.tables[1]
                Third_table= document.tables[2]
                Fourth_table= document.tables[3]

                # new_data = ["Jan-24","3226","241.5","$17,063.92 ","$91,761.25"]

                months_dict = {
                    1: "January",
                    2: "February",
                    3: "March",
                    4: "April",
                    5: "May",
                    6: "June",
                    7: "July",
                    8: "August",
                    9: "September",
                    10: "October",
                    11: "November",
                    12: "December"
                }
                date_dict = {
                    1: "1st",
                    2: "2nd",
                    3: "3rd",
                    4: "4th",
                    5: "5th",
                    6: "6th",
                    7: "7th",
                    8: "8th",
                    9: "9th",
                    10: "10th",
                    11: "11th",
                    12: "12th",
                    13: "13th",
                    14: "14th",
                    15: "15th",
                    16: "16th",
                    17: "17th",
                    18: "18th",
                    19: "19th",
                    20: "20th",
                    21: "21st",
                    22: "22nd",
                    23: "23rd",
                    24: "24th",
                    25: "25th",
                    26: "26th",
                    27: "27th",
                    28: "28th",
                    29: "29th",
                    30: "30th",
                    31: "31st"
                }




                Year= int(data1[0]["First Table"][0]["DATE OF SERVICE PROVIDED"][:4])
                Month= months_dict[round(int(data1[0]["First Table"][0]["DATE OF SERVICE PROVIDED"][5:7]))]
                first_date= date_dict[round(int(data1[0]["First Table"][0]["DATE OF SERVICE PROVIDED"][8:]))]
                last_date= date_dict[round(int(data1[0]["First Table"][-1]["DATE OF SERVICE PROVIDED"][8:]))]

                if data1[0]['Second Table']['Total Regular Hours'].endswith(".0"):
                    data1[0]['Second Table']['Total Regular Hours']= data1[0]['Second Table']['Total Regular Hours'][:-2]

                if data1[0]['Second Table']['Back-up Hours'].endswith(".0"):
                    data1[0]['Second Table']['Back-up Hours']= data1[0]['Second Table']['Back-up Hours'][:-2]

                if data1[0]['Second Table']['Grand Total'].endswith(".0"):
                    data1[0]['Second Table']['Grand Total']= data1[0]['Second Table']['Grand Total'][:-2]


                if data1[0]['Third Table']['PROGRAM MANAGEMENT'].endswith(".0"):
                    data1[0]['Third Table']['PROGRAM MANAGEMENT']= data1[0]['Third Table']['PROGRAM MANAGEMENT'][:-2]

                if data1[0]['Third Table']['ADMINISTRATIVE SUPPORT'].endswith(".0"):
                    data1[0]['Third Table']['ADMINISTRATIVE SUPPORT']= data1[0]['Third Table']['ADMINISTRATIVE SUPPORT'][:-2]

                if data1[0]['Third Table']['Total Labor Hours'].endswith(".0"):
                    data1[0]['Third Table']['Total Labor Hours']= data1[0]['Third Table']['Total Labor Hours'][:-2]


                pas=round(int(data1[0]['Second Table']['Total Regular Hours']))
                backup= round(int(data1[0]['Second Table']['Back-up Hours']))
                total_hours= round(int(data1[0]['Second Table']['Grand Total']))
                pm= round(int(data1[0]['Third Table']['PROGRAM MANAGEMENT']))
                pa= round(int(data1[0]['Third Table']['ADMINISTRATIVE SUPPORT']))
                total_labor_hour= round(int(data1[0]['Third Table']['Total Labor Hours']))


                def add_row_in_last_table(last_table,new_data):

                    def set_cell_border(cell, **kwargs):
                        """
                        Set cell`s border
                        Usage:
                        set_cell_border(
                            cell,
                            top={"sz": 12, "val": "single", "color": "FF0000", "space": "0"},
                            bottom={"sz": 12, "val": "single", "color": "00FF00", "space": "0"},
                            left={"sz": 12, "val": "single", "color": "0000FF", "space": "0"},
                            right={"sz": 12, "val": "single", "color": "000000", "space": "0"},
                        )
                        """
                        tc = cell._element
                        tcPr = tc.get_or_add_tcPr()

                        # Check for each border and set it
                        for border_name, border_attrs in kwargs.items():
                            border = OxmlElement(f"w:{border_name}")
                            for key, value in border_attrs.items():
                                border.set(qn(f"w:{key}"), str(value))
                            tcPr.append(border)


                    new_row = last_table.add_row()

                    last_row_index = len(last_table.rows) - 2  # -2 to skip the newly added row
                    last_row = last_table.rows[last_row_index]

                    for i, cell in enumerate(new_row.cells):
                        last_cell = last_row.cells[i]
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                # Copy font name and size from the last cell
                                last_run = last_cell.paragraphs[0].runs[0]
                                run.font.name = last_run.font.name
                                run.font.size = last_run.font.size

                    for i, cell in enumerate(new_row.cells):
                        cell.text = new_data[i]
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.name = last_cell.paragraphs[0].runs[0].font.name
                                run.font.size = last_cell.paragraphs[0].runs[0].font.size

                    last_row = last_table.rows[-1]  # the newly added row
                    for cell in last_row.cells:
                        for paragraph in cell.paragraphs:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

                    for cell in last_row.cells:
                        set_cell_border(
                            cell,
                            top={"sz": 10, "val": "single", "color": "000000", "space": "0"},
                            bottom={"sz": 10, "val": "single", "color": "000000", "space": "0"},
                            left={"sz": 10, "val": "single", "color": "000000", "space": "0"},
                            right={"sz": 10, "val": "single", "color": "000000", "space": "0"},
                        )




                def add_row_in_table(last_table,new_data):

                    def set_cell_border(cell, **kwargs):
                        """
                        Set cell`s border
                        Usage:
                        set_cell_border(
                            cell,
                            top={"sz": 12, "val": "single", "color": "FF0000", "space": "0"},
                            bottom={"sz": 12, "val": "single", "color": "00FF00", "space": "0"},
                            left={"sz": 12, "val": "single", "color": "0000FF", "space": "0"},
                            right={"sz": 12, "val": "single", "color": "000000", "space": "0"},
                        )
                        """
                        tc = cell._element
                        tcPr = tc.get_or_add_tcPr()

                        # Check for each border and set it
                        for border_name, border_attrs in kwargs.items():
                            border = OxmlElement(f"w:{border_name}")
                            for key, value in border_attrs.items():
                                border.set(qn(f"w:{key}"), str(value))
                            tcPr.append(border)


                    new_row = last_table.add_row()

                    last_row_index = len(last_table.rows) - 2  # -2 to skip the newly added row
                    last_row = last_table.rows[last_row_index]

                    for i, cell in enumerate(new_row.cells):
                        last_cell = last_row.cells[i]
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                # Copy font name and size from the last cell
                                last_run = last_cell.paragraphs[0].runs[0]
                                run.font.name = last_run.font.name
                                run.font.size = last_run.font.size

                    for i, cell in enumerate(new_row.cells):
                        cell.text = new_data[i]
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.name = last_cell.paragraphs[0].runs[0].font.name
                                run.font.size = last_cell.paragraphs[0].runs[0].font.size

                    last_row = last_table.rows[-1]  # the newly added row
                    for cell in last_row.cells:
                        for paragraph in cell.paragraphs:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

                    for cell in last_row.cells:
                        set_cell_border(
                            cell,
                            top={"sz": 10, "val": "single", "color": "000000", "space": "0"},
                            bottom={"sz": 10, "val": "single", "color": "000000", "space": "0"},
                            left={"sz": 10, "val": "single", "color": "000000", "space": "0"},
                            right={"sz": 10, "val": "single", "color": "000000", "space": "0"},
                        )


                def modify_text(document):
                    for p in document.paragraphs:
                        # # print("\n********New paragraph Here *********\n",p.text)
                        if "month of April  2024 – Precision commenced services on April  1st 2024" in p.text:
                            # print(p.text)
                            p.text=f'For the month of {Month} {Year} – Precision commenced services on {Month} {first_date} {Year}. Primary PAS attendant continued government security clearance process.'
                            for run in p.runs:
                                run.font.size = Pt(12)



                def First_table_modification(First_table,Month,first_date,Year,last_date,pas,backup,total_hours):
                    for row in First_table.rows:
                        for cell in row.cells:
                            if "\nApril  1st 2024 thru\nApril 30th   2024\n18  (PAS) + 18 (Backup)=\n36 Hours" in cell.text:
                                # # print("yes")
                                cell.text=""
                                p= cell.add_paragraph()
                                run= p.add_run(f"\n{Month} {first_date} {Year} thru\n{Month} {last_date} {Year}\n{pas}  (PAS) + {backup} (Backup)=\n{total_hours} Hours")
                                run.italic=True

                            if "PAS Hours = 18\nPAS Backup Hours = 18\nTotal Hours= 36\n*Backup = first 2 Hours per shift (9 shifts)" in cell.text:
                                cell.text=""
                                p= cell.add_paragraph()
                                run= p.add_run(f"PAS Hours = {pas}\nPAS Backup Hours = {backup}\nTotal Hours= {total_hours}\n*Backup = first 2 Hours per shift (9 shifts)")
                                run.bold=True
                                run.italic=True
                        
                def Second_table_modification(Second_table,pas,backup,pm,pa):
                    for row in Second_table.rows:
                        f=[]
                        for cell in row.cells:
                            if 'Value1' in cell.text:
                                cell.text=""
                                p= cell.add_paragraph()
                                run= p.add_run(str(pas))
                                run.font.size = Pt(12)
                                if "\n" in cell.text:
                                    cell.text=cell.text.replace("\n","")

                            if 'Value2' in cell.text:
                                cell.text=""
                                p= cell.add_paragraph()
                                run= p.add_run(str(backup))
                                run.font.size = Pt(12)
                                if "\n" in cell.text:
                                    cell.text=cell.text.replace("\n","")

                            if 'Value3' in cell.text:
                                cell.text=""
                                p= cell.add_paragraph()
                                run= p.add_run(str(pm))
                                run.font.size = Pt(12)

                            if 'Value4' in cell.text:
                                cell.text=""
                                p= cell.add_paragraph()
                                run= p.add_run(str(pa))
                                run.font.size = Pt(12)
                                if "\n" in cell.text:
                                    cell.text=cell.text.replace("\n","")
                            f.append(cell.text)
                        # print(f)

                def Third_table_modification(Third_table,data1,pas,backup,total_hours):
                    for row_index,row in enumerate(Third_table.rows):
                        if row_index!=0:
                            for cell_index, cell in enumerate(row.cells):
                                try:
                                    new_cell_text= list(data1[0]["First Table"][row_index-1].values())[cell_index]

                                    if new_cell_text == None:
                                        new_cell_text=  " "
                                    if new_cell_text.endswith(".0"):
                                        new_cell_text=  new_cell_text[:-2]
                                    # # print(cell.text)
                                    
                                    if cell.text.endswith("Value1"):
                                        new_cell_text=  str(pas)

                                    if cell.text.endswith("Value2"):
                                        new_cell_text=  str(backup)

                                    if cell.text.endswith("Value3"):
                                        new_cell_text=  str(total_hours)

                                    cell.text= new_cell_text
                                except Exception as e:
                                    if cell.text.endswith("Value1"):
                                        cell.text=  str(pas)

                                    if cell.text.endswith("Value2"):
                                        cell.text=  str(backup)

                                    if cell.text.endswith("Value3"):
                                        cell.text=  str(total_hours)
                                    # print(e)
                            #     f.append(cell.text)
                            # # print(f)


                def Fourth_table_modification(Fourth_table,pas,backup,pm,pa,total_labor_hour):
                    total_labor_hour= "            "+str(total_labor_hour)
                    for row in Fourth_table.rows:
                        for cell in row.cells:
                            if "Value1" in cell.text:
                                cell.text=  "             "+ str(pas)

                            if "Value2" in cell.text:
                                cell.text=  "             "+str(backup)

                            if "Value3" in cell.text:
                                cell.text=  "\n             "+str(pm)

                            if "Value4" in cell.text:
                                cell.text=  "\n             "+str(pa)

                            if "Value5" in cell.text:
                                cell.text=""
                                p= cell.add_paragraph()
                                run= p.add_run(total_labor_hour)
                                run.bold=True
                                # if "\n" in cell.text:
                                #     cell.text=cell.text.replace("\n","")

                def last_table_modification(Last_table,data2):

                    for dic in data2:
                        new_data= list(dic.values())
                        for i_val, val in enumerate(new_data):
                            if val == None:
                                new_data[i_val]= " "

                        add_row_in_table(Last_table,new_data)

                    # Temp_list=[]
                    # for row_index,row in enumerate(Last_table.rows):
                    #     if row_index > 1:
                    #         Temp_list2=[]
                    #         for cell_index, cell in enumerate(row.cells):
                    #             Temp_list2.append(cell.text)
                    #         Temp_list.append(Temp_list2)
                    # # print(Temp_list)


                modify_text(document)
                First_table_modification(First_table,Month,first_date,Year,last_date,pas,backup,total_hours)
                Second_table_modification(Second_table,pas,backup,pm,pa)
                Third_table_modification(Third_table,data1,pas,backup,total_hours)
                Fourth_table_modification(Fourth_table,pas,backup,pm,pa,total_labor_hour)
                last_table_modification(last_table,data2)


                # Save the document to a BytesIO object
                docx_buffer = BytesIO()
                document.save(docx_buffer)
                docx_buffer.seek(0)


                # with col1:
                #     if st.button("View Report Here"):
                #         pdf_path = "document.pdf"

                #         # Read the PDF file
                #         with open(pdf_path, "rb") as f:
                #             pdf_bytes = f.read()
                #         pdf_display = f'<iframe src="data:application/pdf;base64,{pdf_bytes.encode("base64")}#toolbar=0" width="700" height="900" type="application/pdf"></iframe>'
                #         st.markdown(pdf_display, unsafe_allow_html=True)


                # with col2:
                
                st.write("Report generated successfully")
                st.download_button(
                    label="Download DOCX",
                    data=docx_buffer,
                    file_name="output.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
                # st.write("Report downloaded successfully")
                break
            except Exception as e:
                print(e)
                pass 
                








